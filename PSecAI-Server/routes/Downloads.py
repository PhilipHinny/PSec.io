from flask import send_file, Blueprint, jsonify
from io import BytesIO
from docx import Document  
from database import get_db_connection

Download_bp = Blueprint("Downloaded", __name__)

@Download_bp.route("/download_report/<filename>", methods=["GET"])
def download_report(filename):
    try:
        
        db = get_db_connection()
        generate_collection = db["Generated_Reports"]

        filename = filename.replace('%20', ' ')  

        report = generate_collection.find_one({"filename": filename})

        if report is None:
            return jsonify({"error": "File not found"}), 404

        report_text = report.get("report_text")
        if not report_text:
            return jsonify({"error": "Report text is missing"}), 500

        doc = Document()
        doc.add_heading(filename, 0)  
        doc.add_paragraph(report_text)  

      
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)  

        return send_file(doc_io, as_attachment=True, download_name=f"{filename}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        print(f"Error downloading file: {e}")
        return jsonify({"error": "Failed to download file"}), 500
